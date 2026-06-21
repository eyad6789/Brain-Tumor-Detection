# -*- coding: utf-8 -*-
"""Render report-en.docx and report-ar.docx from doc_data.SECTIONS.

Editable Microsoft Word versions of the deep-dive. The Arabic edition sets
paragraph bidi + run rtl + bidiVisual tables and an Arabic complex-script font,
so Word shapes Arabic correctly and lays the document out right-to-left.

Run:  ../.venv/bin/python build_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import doc_data
from i18n import tr
import theme as T

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "..", "assets")
OUT = os.path.join(HERE, "..")

# Print-friendly light palette (Word) ------------------------------------
INK = RGBColor(0x1A, 0x24, 0x33)
INKS = RGBColor(0x0B, 0x14, 0x22)
MUT = RGBColor(0x5D, 0x6B, 0x80)
PRIMARY = RGBColor(0x08, 0x91, 0xA8)
ACCENT = RGBColor(0x5B, 0x54, 0xE6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CODEINK = RGBColor(0x16, 0x20, 0x33)

SHADE_HEAD = "DCE6F2"
SHADE_CODE = "F1F5FB"
SHADE_KVK = "EEF3FA"
CALL_BG = {"info": "EAF7FA", "warn": "FBF1E2", "good": "EAF6F0"}
CALL_BAR = {"info": "0891A8", "warn": "B9701A", "good": "0F8A55"}

META = {
    "en": {"title": "NeuroScan — Technical Deep-Dive",
           "sub": "Explainable Brain-Tumor MRI Classification · ResNet50 → GradCAM → Local-LLM Report",
           "kicker": "COMPLETE TECHNICAL REFERENCE", "toc": "Contents",
           "disc": "Educational & research prototype — not a medical device. Never use for diagnosis."},
    "ar": {"title": "NeuroScan — تقرير تقني معمّق",
           "sub": "تصنيف قابل للتفسير لأورام الدماغ · ResNet50 ثم GradCAM ثم تقرير بنموذج لغوي محلي",
           "kicker": "مرجع تقني كامل", "toc": "المحتويات",
           "disc": "نموذج تعليمي وبحثي — ليس جهازًا طبيًا. لا تستخدمه للتشخيص."},
}
CALL_LABEL = {"info": {"en": "Note", "ar": "ملاحظة"},
              "warn": {"en": "Important", "ar": "تنبيه مهم"},
              "good": {"en": "Good to know", "ar": "جيد أن تعرف"}}


def _el(tag):
    return OxmlElement(tag)


def set_para_rtl(p):
    pPr = p._p.get_or_add_pPr()
    b = _el("w:bidi"); pPr.append(b)


def set_run_rtl(r):
    rPr = r._element.get_or_add_rPr()
    rtl = _el("w:rtl"); rtl.set(qn("w:val"), "true"); rPr.append(rtl)


def set_cs_font(r, name):
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = _el("w:rFonts"); rPr.append(rf)
    rf.set(qn("w:cs"), name)
    rf.set(qn("w:ascii"), T.FONT_DISPLAY)
    rf.set(qn("w:hAnsi"), T.FONT_DISPLAY)


def shade(element, fill):
    """Apply background shading to a paragraph (_p) or table cell (_tc)."""
    pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shd = _el("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pr.append(shd)


def para_border(p, *, bottom=None, left=None, size=8):
    pPr = p._p.get_or_add_pPr()
    bdr = _el("w:pBdr")
    if bottom:
        b = _el("w:bottom"); b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "6"); b.set(qn("w:color"), bottom); bdr.append(b)
    if left:
        l = _el("w:left"); l.set(qn("w:val"), "single"); l.set(qn("w:sz"), str(size))
        l.set(qn("w:space"), "8"); l.set(qn("w:color"), left); bdr.append(l)
    pPr.append(bdr)


def table_bidi(table):
    tblPr = table._tbl.tblPr
    bv = _el("w:bidiVisual"); tblPr.append(bv)


class DocBuilder:
    def __init__(self, lang):
        self.lang = lang
        self.rtl = lang == "ar"
        self.arface = T.FONT_AR
        self.doc = Document()
        self._base_style()

    def _base_style(self):
        st = self.doc.styles["Normal"]
        st.font.name = T.FONT_DISPLAY
        st.font.size = Pt(11)
        st.font.color.rgb = INK
        for sec in self.doc.sections:
            sec.left_margin = sec.right_margin = Inches(0.9)
            sec.top_margin = sec.bottom_margin = Inches(0.85)
            if self.rtl:
                sectPr = sec._sectPr
                b = _el("w:bidi"); sectPr.append(b)

    # ── run / paragraph helpers ──────────────────────────────────────────
    def run(self, p, text, *, size=11, color=INK, bold=False, mono=False):
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
        r.font.name = T.FONT_MONO if mono else T.FONT_DISPLAY
        if self.rtl and not mono:
            set_run_rtl(r); set_cs_font(r, self.arface)
        return r

    def para(self, text="", *, size=11, color=INK, bold=False, style=None,
             align=None, before=4, after=4, mono=False):
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        elif self.rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if self.rtl:
            set_para_rtl(p)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if text:
            self.run(p, text, size=size, color=color, bold=bold, mono=mono)
        return p

    # ── blocks ────────────────────────────────────────────────────────────
    def block(self, b):
        t = b["type"]
        if t == "p":
            self.para(tr(b["text"], self.lang))
        elif t == "h":
            p = self.para(before=12, after=2)
            self.run(p, tr(b["text"], self.lang), size=13.5, color=INKS, bold=True)
        elif t == "bullets":
            for it in b["items"]:
                self.para(tr(it, self.lang),
                          style="List Bullet" + (" " if False else ""))
        elif t == "num":
            for it in b["items"]:
                self.para(tr(it, self.lang), style="List Number")
        elif t == "table":
            self.table([[tr(c, self.lang) for c in b["head"]]],
                       [[tr(c, self.lang) for c in r] for r in b["rows"]])
        elif t == "code":
            self.code(b["code"])
        elif t == "callout":
            self.callout(b)
        elif t == "image":
            self.image(b)
        elif t == "kv":
            self.kv(b)

    def table(self, head_rows, body_rows):
        ncol = len(head_rows[0])
        tbl = self.doc.add_table(rows=0, cols=ncol)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        if self.rtl:
            table_bidi(tbl)
        for hr in head_rows:
            cells = tbl.add_row().cells
            for j, txt in enumerate(hr):
                shade(cells[j]._tc, SHADE_HEAD)
                self._cell(cells[j], txt, bold=True, color=INKS)
        for r in body_rows:
            cells = tbl.add_row().cells
            for j, txt in enumerate(r):
                self._cell(cells[j], txt, bold=(j == 0), color=INK)
        self.para(before=2, after=6)

    def _cell(self, cell, text, *, bold, color, size=10.5, mono=False):
        cell.vertical_alignment = 1
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        if self.rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_para_rtl(p)
        self.run(p, text, size=size, color=color, bold=bold, mono=mono)

    def code(self, codetext):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        shade(cell._tc, SHADE_CODE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        lines = codetext.split("\n")
        for i, line in enumerate(lines):
            r = p.add_run(line)
            r.font.name = T.FONT_MONO; r.font.size = Pt(9.5); r.font.color.rgb = CODEINK
            if i < len(lines) - 1:
                r.add_break()
        self.para(before=2, after=6)

    def callout(self, b):
        tone = b["tone"]
        tbl = self.doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        shade(cell._tc, CALL_BG[tone])
        # colored left bar via cell border
        tcPr = cell._tc.get_or_add_tcPr()
        borders = _el("w:tcBorders")
        side = "w:end" if self.rtl else "w:start"
        bar = _el(side); bar.set(qn("w:val"), "single"); bar.set(qn("w:sz"), "24")
        bar.set(qn("w:space"), "0"); bar.set(qn("w:color"), CALL_BAR[tone])
        borders.append(bar); tcPr.append(borders)
        ph = cell.paragraphs[0]
        ph.paragraph_format.space_before = Pt(3); ph.paragraph_format.space_after = Pt(1)
        if self.rtl:
            ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_para_rtl(ph)
        lab = CALL_LABEL[tone][self.lang]
        self.run(ph, f"{lab} · {tr(b['title'], self.lang)}", size=11.5, color=INKS, bold=True)
        pb = cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(1); pb.paragraph_format.space_after = Pt(3)
        if self.rtl:
            pb.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_para_rtl(pb)
        self.run(pb, tr(b["text"], self.lang), size=11, color=INK)
        self.para(before=2, after=6)

    def image(self, b):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join(ASSETS, b["src"]), width=Inches(6.2))
        if b.get("caption"):
            cap = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8)
            self.run(cap, tr(b["caption"], self.lang), size=9.5, color=MUT)

    def kv(self, b):
        tbl = self.doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        if self.rtl:
            table_bidi(tbl)
        for k, v in b["pairs"]:
            cells = tbl.add_row().cells
            shade(cells[0]._tc, SHADE_KVK)
            self._cell(cells[0], tr(k, self.lang), bold=True, color=INKS, size=10.5)
            self._cell(cells[1], tr(v, self.lang), bold=False, color=INK, size=10.5)
        self.para(before=2, after=6)

    # ── document scaffold ────────────────────────────────────────────────
    def cover(self):
        m = META[self.lang]
        self.para(before=40, after=0)
        k = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
        self.run(k, m["kicker"], size=11, color=PRIMARY, bold=True)
        t = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4)
        self.run(t, m["title"], size=30, color=INKS, bold=True)
        s = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=10)
        self.run(s, m["sub"], size=12.5, color=MUT)
        rule = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
        para_border(rule, bottom="0891A8", size=18)

    def toc(self):
        m = META[self.lang]
        h = self.para(before=24, after=6)
        self.run(h, m["toc"], size=18, color=INKS, bold=True)
        para_border(h, bottom="DCE6F2", size=10)
        for s in doc_data.SECTIONS:
            p = self.para(before=1, after=1)
            self.run(p, f"{s['num']}   ", size=11, color=PRIMARY, bold=True)
            self.run(p, tr(s["title"], self.lang), size=11, color=INK)
        self.doc.add_page_break()

    def section(self, s):
        head = self.para(before=16, after=4)
        self.run(head, f"{s['num']}   ", size=20, color=PRIMARY, bold=True)
        self.run(head, tr(s["title"], self.lang), size=20, color=INKS, bold=True)
        para_border(head, bottom="DCE6F2", size=12)
        if s.get("lead"):
            lead = self.para(before=4, after=6)
            self.run(lead, tr(s["lead"], self.lang), size=12.5, color=MUT)
        for b in s["blocks"]:
            self.block(b)

    def build(self):
        self.cover()
        self.toc()
        for s in doc_data.SECTIONS:
            self.section(s)
        # footer disclaimer
        f = self.para(before=18, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        para_border(f, bottom="DCE6F2", size=6)
        f2 = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=0)
        self.run(f2, META[self.lang]["disc"], size=9.5, color=MUT)
        return self.doc


def main():
    for lang in ("en", "ar"):
        b = DocBuilder(lang)
        doc = b.build()
        path = os.path.join(OUT, f"report-{lang}.docx")
        doc.save(path)
        print(f"wrote report-{lang}.docx  ({len(doc.paragraphs)} paragraphs)")


if __name__ == "__main__":
    main()
